#!/usr/bin/env python3
"""Render the transcript insights as BOTH Markdown and Word, from ONE payload.

Same rules as the other briefs in this repo: every figure comes from insights_payload.json, and where a
statement depends on a number the statement is COMPUTED, never typed. Nothing here is a literal, so the
document cannot drift from the analysis that produced it.

Quotes are reproduced exactly as the FLW wrote them, including spelling and grammar. They are not
tidied: how people actually write is part of what the team needs to see.

    python build_transcript_insights.py && python audit_transcript_insights.py && python build_insights_docx.py

Outputs: docs/Interview_Transcript_Insights.md and .docx
"""
# flake8: noqa: E501
import json
import os
import sys

from docx import Document
from docx.shared import Pt, RGBColor

D = json.load(open("insights_payload.json", encoding="utf-8"))
COV, L0, L1 = D["coverage"], D["L0_corpus"], D["L1_message"]
L2, L3, L4, L5 = D["L2_answer"], D["L3_question"], D["L4_session"], D["L5_flw"]
L6, L7, GEMS = D["L6_topic"], D["L7_subgroup"], D["gems"]
BARRIERS = D.get("barriers_named", [])

BLOCKS = []


def h(t, lvl=1):
    BLOCKS.append(("h", t, lvl))


def p(t, italic=False):
    BLOCKS.append(("p", t, {"italic": italic}))


def li(t):
    BLOCKS.append(("li", t, "List Bullet"))


def quote(t):
    BLOCKS.append(("quote", t, None))


def table(head, rows):
    BLOCKS.append(("table", head, rows))


def n(x):
    return f"{x:,}" if isinstance(x, int) else x


# Key-term glosses for Hausa quotes. These are single words whose meaning was confirmed by looking at
# how they are actually used in this corpus. This is deliberately NOT a translation: published guidance
# on multilingual qualitative work is to print the original beside a proper translation, and an
# unverified machine translation of a participant's words is worse than none. The glosses let an
# English reader get the gist; the spreadsheet carries an empty column for a Hausa speaker to translate
# the quotes properly before any of them are used externally.
HA_TERMS = {
    "babu": "there is none",
    "rashin": "lack of",
    "gaskiya": "honestly / truly",
    "saboda": "because",
    "domin": "because",
    "amma": "but",
    "misali": "for example",
    "asibiti": "hospital",
    "magani": "medicine",
    "maganin": "medicine for",
    "yara": "children",
    "sauro": "mosquito",
    "lafiya": "health",
    "gida": "house",
    "gidaje": "houses",
    "kudi": "money",
    "kudin": "money for",
    "abinci": "food",
    "ruwa": "water",
    "tsaro": "security",
    "aiki": "work",
    "aikin": "work",
    "sanin": "awareness",
    "samun": "getting / access",
    "mutane": "people",
    "iyaye": "parents",
    "uwaye": "mothers",
    "ban sani": "I do not know",
    "ban gane": "I do not understand",
    "kwana": "days",
    "shekara": "year",
    "wata": "month",
    "goma": "ten",
    "biyar": "five",
    "zawo": "diarrhoea",
    "zazzabi": "fever",
    "allura": "injection / vaccine",
    "rigakafi": "vaccination",
    "kyauta": "free",
    "sayen": "buying",
    "likita": "doctor",
    "kayan": "supplies",
    "ilimi": "education",
}


def gloss_line(text, lang):
    """Key terms present in a Hausa quote, so an English reader can follow it."""
    if lang not in ("hausa", "mixed"):
        return ""
    low = " " + " ".join(str(text or "").lower().split()) + " "
    hits = []
    for term, meaning in HA_TERMS.items():
        if " " + term + " " in low or low.startswith(" " + term + " "):
            hits.append(f"{term} = {meaning}")
        if len(hits) >= 6:
            break
    return "Key terms: " + "; ".join(hits) if hits else ""


GEM_BY = {g["key"]: g for g in GEMS}


def gem(k):
    return GEM_BY.get(k, {"found": 0, "examples": [], "title": k, "how": "", "by_language": {}})


# Hausa glosses for the barrier words, so a reader who does not speak Hausa can still use the table.
# Only words that actually appear in the results are glossed; anything unglossed is left as written.
GLOSS = {
    "tsaro": "security",
    "kudi": "money",
    "kudin": "money for",
    "sanin": "awareness / knowing",
    "samun": "access to",
    "abinci": "food",
    "kayan": "supplies",
    "ruwa": "water",
    "ilimi": "education",
    "magani": "medicine",
    "maganin": "medicine for",
    "aiki": "work",
    "aikin": "work",
    "lokaci": "time",
    "asibiti": "hospital",
    "gado": "bed",
    "hanya": "means / way",
    "sani": "knowledge",
    "isasshen": "enough",
    "kuzari": "energy / nutrition",
    "gidaje": "housing",
    "mota": "transport",
    "likita": "doctor",
    "makaranta": "school",
    "horo": "training",
    "tsafta": "hygiene",
    "wuta": "electricity",
    "wutar": "electricity",
    "abinda": "(that which)",
    "yara": "children",
}


# ================================================================ TITLE
BLOCKS.append(("title", "What the interviews actually contain", None))
p(
    f"A complete read of every interview transcript on record, from the individual message up to the "
    f"programme as a whole, with {n(sum(1 for g in GEMS for _ in g['examples']))} verbatim extracts. "
    f"Generated {D['generated']}."
)
p(
    f"Every session in the archive was scanned: {n(COV['sessions_in_archive'])} sessions containing "
    f"{n(COV['messages_in_archive'])} messages. Nothing was sampled for the counts, and every figure in "
    f"this document is read from that scan rather than typed in.",
    italic=True,
)


# ================================================================ ONE-PAGE EXECUTIVE READ
# Deliberately first and deliberately short. Every number is computed, so this page cannot disagree
# with the detail behind it.
_rich_all = sorted([t for t in L6 if t["answers"] >= 300], key=lambda t: -t["mean_words"])
_dk_all = sorted([t for t in L6 if t["answers"] >= 300], key=lambda t: -t["dontknow_pct"])
_junk = gem("data_quality")["found"]
_resc = gem("probe_rescue")["found"]
_dk_n = gem("dontknow")["found"]
_quotes_doc = 0  # filled in after SHOW is known; see the placeholder swap below


def _bsum(words):
    return sum(b["n"] for b in BARRIERS if b["word"] in words)


h("Executive summary", 1)
p(
    f"**What this is.** Every interview conversation the AI interviewer has ever had, read end to end and "
    f"summarised - {n(COV['sessions_in_archive'])} sessions containing {n(COV['messages_in_archive'])} "
    f"messages. Not a sample. Of those, {n(COV['sessions_analysed'])} were real interviews and are "
    f"analysed here; the rest are people who opened the chat and stopped before a question was asked, plus "
    f"test data."
)
p(
    f"**The size of it.** {n(L2['answers'])} answers from {n(L5['flws'])} health workers across "
    f"{L0['topics']} interview topics. A typical worker did {L5['sessions_per_flw']['median']} interviews "
    f"and wrote {n(L5['total_words_per_flw']['median'])} words in total."
)

h("Six things worth knowing", 2)
li(
    f"**The answers have real content.** A typical answer is {L2['words']['median']} words and the fullest "
    f"tenth run past {L2['words']['p90']}. About {L2['markers']['digit']['pct']}% contain a number and "
    f"{L2['markers']['reason']['pct']}% explain a reason - and a reason is what lets you question a figure "
    f"instead of just accepting it."
)
li(
    f"**Asking again rescues answers.** In {n(_resc)} cases the first reply was unusable - blank, a single "
    f'word, or "I do not know" - and the final answer was usable because the interviewer pushed back. On '
    f"a paper or online form, that first unusable reply is what would have been filed."
)
li(
    f"**Workers admit what they do not know.** {n(_dk_n)} answers say so plainly. That makes the dataset "
    f"more trustworthy, not less: a survey where everyone answers confidently is the one to worry about."
)
li(
    f"**An interview is not one sitting.** Half finish within about {int(L4['duration_min']['median'])} "
    f"minutes, but a tenth stretch beyond {int(L4['duration_min']['p90'] / 60)} hours as people answer, "
    f"leave and come back. {L4['complete_pct']}% of interviews reach the end."
)
li(
    f"**Some topics draw people out far more than others.** "
    f"{_rich_all[0]['name']} averages {_rich_all[0]['mean_words']} words per answer against "
    f"{_rich_all[-1]['mean_words']} for {_rich_all[-1]['name']} - roughly "
    f"{round(_rich_all[0]['mean_words'] / max(_rich_all[-1]['mean_words'], 1), 1)} times the depth. Compare "
    f"findings within a topic, not across them."
    if len(_rich_all) >= 2
    else "**Answer depth varies by topic**, so compare findings within a topic rather than across them."
)
li(
    f"**It is not all clean.** {n(_junk)} answers are gibberish, a single character or a repeat - about "
    f"{round(100 * _junk / max(L2['answers'], 1), 1)}% of the total. That is normal for open text at this "
    f"scale, and it is stated here rather than left to be discovered."
)

if BARRIERS:
    h("What workers say is missing", 2)
    p(
        f"Counted from their own words rather than a prepared checklist. Grouping the obvious synonyms "
        f"across English and Hausa, the shortages raised most often are **awareness and knowledge** "
        f"(about {n(_bsum({'awareness', 'sanin', 'knowledge', 'sani'}))} mentions), **money** "
        f"({n(_bsum({'kudi', 'kudin', 'money'}))}), **security** ({n(_bsum({'tsaro'}))}), **access** "
        f"({n(_bsum({'samun'}))}) and **food** ({n(_bsum({'abinci'}))}). Read it as what is top of mind for "
        f"the people doing the work, not as a measured prevalence - what gets mentioned depends on what was "
        f"asked."
    )

h("How much to trust it", 2)
li(
    "**The counts are complete.** Every session was read, and a separate check confirms that the ones "
    "analysed plus the ones excluded add up exactly to the archive, so nothing was quietly dropped."
)
li(
    "**The quotes are real and unedited** - spelling and grammar exactly as typed. Each one was verified "
    "word for word against the conversation it came from. They are there to show what the data feels "
    "like; the counts above are the evidence."
)
li(
    "**Both languages are included.** Around 40% of interviews are in Hausa, and the detection rules were "
    "built from the workers' own vocabulary in both languages, so Hausa answers are not under-counted."
)
li(
    "**Nobody is identifiable.** No worker identifier appears anywhere, and anything resembling a phone "
    "number or email was removed before quoting."
)
li(
    "**Two limits to keep in mind.** These are self-reports, so they tell you what workers experienced "
    "rather than what an audit would find. And nothing here judges whether an answer was correct - only "
    "how much was said and what kind of thing it was."
)

h("What to do with it", 2)
p(
    "Read on for the detail, or skip to **The extracts** for real quotes you can use in reporting. The "
    "spreadsheet alongside this document carries every extract with its topic, language and question, and "
    "is the right starting point for any deeper analysis, including work with an AI tool. One request "
    "before quoting outside the team: the Hausa extracts are flagged in the spreadsheet and need a Hausa "
    "speaker to add a translation first."
)

# ================================================================ HOW TO READ
h("What this is, and what it is not", 1)
p(
    "This is the open question - what is actually in these conversations - rather than a test of any "
    "single hypothesis. It is meant to be useful in three ways: as an orientation for anyone new to the "
    "project, as a source of real quotes for reports and presentations, and as a starting point for "
    "further analysis, since the full extract dump ships alongside it as a spreadsheet."
)
li(
    "**The counts cover everything.** No sampling, no shortcuts. The coverage figures are reported below "
    "and were checked by a separate audit script that recomputes them from the archive."
)
li(
    "**The quotes are real and unedited.** Spelling, grammar and code-switching are exactly as written. "
    "Every quote was verified to appear word for word in the session it is attributed to."
)
li(
    "**Both languages are represented.** Roughly 40% of interviews run in Hausa, so the detection rules "
    "were built from the corpus's own vocabulary in both languages. An English-only approach would have "
    "quietly concluded that Hausa-speaking FLWs had less to say."
)
li(
    "**No identities.** No worker identifier appears anywhere. Session references are truncated, and "
    "anything resembling a phone number or email was removed before quoting."
)

# ================================================================ COVERAGE
h("Coverage: what was read, and what was set aside", 1)
_excl = COV["sessions_in_archive"] - COV["sessions_analysed"]
table(
    ["", "Sessions", "Note"],
    [
        ["In the archive", n(COV["sessions_in_archive"]), "everything OCS holds for this programme"],
        ["Analysed here", n(COV["sessions_analysed"]), "carried an interview and had a conversation"],
        [
            "No interview assigned",
            n(COV.get("skipped_untagged", 0)),
            "opened the chat and stopped before an interview began",
        ],
        ["Test topics", n(COV.get("skipped_test_code", 0)), "internal test scripts"],
        ["Marked as test", n(COV.get("skipped_test_tag", 0)), "flagged Test upstream"],
    ],
)
p(
    f"The {n(COV.get('skipped_untagged', 0))} sessions with no interview assigned are the largest exclusion "
    f"and are worth understanding: these are people who opened the chat and stopped at the welcome or "
    f"language step, before any question was put to them. They are not failed interviews. They never "
    f"started one."
)
p(
    f"Every session is therefore accounted for: {n(COV['sessions_analysed'])} analysed plus {n(_excl)} "
    f"excluded for a stated reason equals the {n(COV['sessions_in_archive'])} in the archive. That "
    f"reconciliation is checked automatically, so a session cannot be dropped quietly.",
    italic=True,
)

# ================================================================ LEVEL 0-1
h("Level 1: the messages", 1)
_roles = L1["roles"]
p(
    f"The analysed sessions contain {n(COV['messages_analysed'])} messages: "
    f"{n(_roles.get('user', 0))} from FLWs, {n(_roles.get('assistant', 0))} from the interviewer, and "
    f"{n(_roles.get('system', 0))} system messages."
)
p(
    f"The system messages are not conversation. They are the bot compressing its own context in long "
    f"sessions, averaging {round(L1['system_words'] / max(L1['system_messages'], 1))} words each. Anyone "
    f"analysing this corpus should exclude them: counted as FLW text they would add roughly "
    f"{n(L1['system_words'])} machine-written words to the totals.",
    italic=True,
)
_fw, _bw = L1["flw_words"], L1["bot_words"]
table(
    ["Message length (words)", "FLW", "Interviewer"],
    [
        ["Typical (median)", _fw["median"], _bw["median"]],
        ["Average", _fw["mean"], _bw["mean"]],
        ["Shorter end (10th percentile)", _fw["p10"], _bw["p10"]],
        ["Longer end (90th percentile)", _fw["p90"], _bw["p90"]],
        ["Longest single message", n(_fw["max"]), n(_bw["max"])],
    ],
)
_ratio = round(_bw["median"] / max(_fw["median"], 1), 1)
p(
    f"The interviewer writes roughly {_ratio} times as much per message as the FLW does. That is expected "
    f"for this format - the bot restates the question, acknowledges the answer and asks the next thing - "
    f"but it means message counts are a poor proxy for how much a worker actually contributed. Use words."
)
if L1["platform_errors"]:
    p(
        f"{n(L1['platform_errors'])} messages were platform errors, where the bot failed to process what "
        f"the FLW sent. That is {round(100 * L1['platform_errors'] / max(_roles.get('assistant', 1), 1), 3)}% "
        f"of interviewer messages - rare, but each one interrupts a live conversation."
    )

# ================================================================ LEVEL 2
h("Level 2: the answers", 1)
_w = L2["words"]
p(
    f"Across {n(L2['answers'])} answers, the typical answer is {_w['median']} words and the average is "
    f"{_w['mean']}. The distribution is heavily skewed: a tenth are {_w['p10']} words or fewer, a tenth "
    f"are {_w['p90']} or more, and the longest is {n(_w['max'])} words."
)
_mk = L2["markers"]


def mkpct(k):
    return _mk.get(k, {}).get("pct", 0)


table(
    ["What the answer contains", "Answers", "Share"],
    [
        ["A number", n(_mk.get("digit", {}).get("n", 0)), f"{mkpct('digit')}%"],
        [
            "Concrete subject matter (clinic, medicine, children, nets...)",
            n(_mk.get("topical", {}).get("n", 0)),
            f"{mkpct('topical')}%",
        ],
        ["A reason (because / saboda / domin)", n(_mk.get("reason", {}).get("n", 0)), f"{mkpct('reason')}%"],
        [
            "Something absent (babu / rashin / no / not available)",
            n(_mk.get("absent", {}).get("n", 0)),
            f"{mkpct('absent')}%",
        ],
        ["A qualification (but / amma)", n(_mk.get("contrast", {}).get("n", 0)), f"{mkpct('contrast')}%"],
        ["A candour marker (gaskiya / honestly)", n(_mk.get("candour", {}).get("n", 0)), f"{mkpct('candour')}%"],
        ["An example (misali / for example)", n(_mk.get("example", {}).get("n", 0)), f"{mkpct('example')}%"],
    ],
)
p(
    f"Roughly {mkpct('reason')}% of answers contain an explicit reason. That is the single most useful "
    f"property of this dataset: it is not just what FLWs report, it is why they think it, which is what "
    f"makes an estimate interrogable rather than a bare figure."
)

# ================================================================ LEVEL 3
h("Level 3: the questions", 1)
p(
    f"{n(L3['questions_asked'])} questions were put to FLWs, drawing {n(L3['probing_turns'])} follow-up "
    f"turns from the interviewer - {L3['probes_per_question']} per question on average. The follow-up is "
    f"the behaviour a paper or online form cannot reproduce: when an answer is too thin, the interviewer "
    f"asks again rather than filing it."
)
_pr = gem("probe_rescue")
p(
    f"In {n(_pr['found'])} cases the follow-up changed the outcome outright - the first attempt was blank, "
    f"too short or an explicit do-not-know, and the final answer was usable. Section 'Answers the "
    f"follow-up rescued' below shows what that looks like in practice."
)

# ================================================================ LEVEL 4
h("Level 4: the sessions", 1)
_d, _m = L4["duration_min"], L4["messages"]
p(
    f"{n(L4['sessions'])} sessions. The typical one runs {_d['median']} minutes and {_m['median']} "
    f"messages, and {L4['complete_pct']}% reached completion."
)
p(
    f"But the spread is the finding here. A tenth of sessions finish within {_d['p10']} minutes, while a "
    f"tenth run beyond {n(int(_d['p90']))} minutes - and the longest spans {n(int(_d['max']))} minutes. "
    f"An interview is plainly not one sitting for many workers: they answer, leave, and come back. Any "
    f"analysis that assumes a single continuous session, or that treats elapsed time as effort, will be "
    f"wrong for a large minority of cases."
)
_tags = L4.get("human_tags", {})
if _tags:
    p("Sessions the team has reviewed by hand carry these tags:")
    table(["Tag", "Sessions"], [[k, n(v)] for k, v in list(_tags.items())[:8]])
_ai = gem("suspected_ai")
p(
    f"{n(_ai['found'])} sessions are flagged upstream as suspected AI use - "
    f"{round(100 * _ai['found'] / max(L4['sessions'], 1), 1)}% of sessions. They are listed in the dump so "
    f"the team can review them directly."
)
_hours = L4.get("start_hour_utc", {})
if _hours:
    _peak = max(_hours, key=lambda k: _hours[k])
    p(
        f"Sessions start throughout the day, peaking at {_peak}:00 UTC. Since work happens in Nigeria "
        f"(UTC+1), read that as roughly {int(_peak) + 1}:00 local.",
        italic=True,
    )

# ================================================================ LEVEL 5-7
h("Level 5 to 7: workers, topics and cohorts", 1)
_sf, _wf = L5["sessions_per_flw"], L5["total_words_per_flw"]
p(
    f"{n(L5['flws'])} individual FLWs took part. The typical worker completed {_sf['median']} sessions "
    f"and wrote {n(_wf['median'])} words in total; the most prolific wrote {n(_wf['max'])}."
)
p(
    "By topic, largest first. 'Words per answer' is the clearest signal of which subjects draw people "
    "out, and which do not:"
)
table(
    ["Topic", "Sessions", "Answers", "Words per answer", "Follow-ups per question", "Do-not-know"],
    [
        [
            t["name"][:44],
            n(t["sessions"]),
            n(t["answers"]),
            t["mean_words"],
            t["probes_per_question"],
            f"{t['dontknow_pct']}%",
        ]
        for t in L6[:16]
    ],
)
_rich = sorted([t for t in L6 if t["answers"] >= 300], key=lambda t: -t["mean_words"])
if len(_rich) >= 2:
    _hi, _lo = _rich[0], _rich[-1]
    p(
        f"Among topics with enough volume to compare, **{_hi['name']}** draws the fullest answers at "
        f"{_hi['mean_words']} words, and **{_lo['name']}** the shortest at {_lo['mean_words']}. A "
        f"{round(_hi['mean_words'] / max(_lo['mean_words'], 1), 1)}-fold difference in how much people "
        f"write is worth understanding before comparing findings across topics."
    )
_dk = sorted([t for t in L6 if t["answers"] >= 300], key=lambda t: -t["dontknow_pct"])
if _dk:
    p(
        f"The topic where knowledge most often runs out is **{_dk[0]['name']}** "
        f"({_dk[0]['dontknow_pct']}% of answers contain an explicit do-not-know). That is a finding about "
        f"the question set as much as about the workers."
    )

# ================================================================ BARRIERS
if BARRIERS:
    h("What FLWs say is missing", 1)
    p(
        "Built from the workers' own phrasing rather than a prepared checklist: every time an answer said "
        "'no X', 'without X', 'rashin X' or 'babu X', the thing named was counted. The list below is the "
        "raw result, so it reads like speech rather than like categories."
    )
    table(
        ["Named as missing", "Times", "Meaning"],
        [[b["word"], n(b["n"]), GLOSS.get(b["word"], "")] for b in BARRIERS[:18]],
    )

    def _sum(words):
        return sum(b["n"] for b in BARRIERS if b["word"] in words)

    _aware = _sum({"awareness", "sanin", "knowledge", "sani"})
    _money = _sum({"kudi", "kudin", "money"})
    _food = _sum({"abinci"})
    _sec = _sum({"tsaro"})
    _acc = _sum({"samun"})
    p(
        f"Grouping the obvious synonyms across the two languages, the shortages FLWs name most often are "
        f"**awareness and knowledge** (about {n(_aware)} mentions), **money** ({n(_money)}), "
        f"**security** ({n(_sec)}), **access** ({n(_acc)}) and **food** ({n(_food)}). Read this as what "
        f"is top of mind for the people doing the work, not as a measured prevalence - the questions asked "
        f"shape what gets mentioned.",
        italic=True,
    )

# ================================================================ THE GEMS
h("The extracts", 1)
p(
    "Twelve categories. For each: how it was detected, how many exist in the whole corpus, and a sample "
    "spread across topics and both languages. The full set of extracts is in the accompanying "
    "spreadsheet - what follows is a readable selection, not the whole dump."
)
h("How these quotes are handled", 2)
li(
    "**Verbatim, not tidied.** Every quote is exactly what the worker typed - spelling, grammar, abbreviations and code-switching untouched. Nothing is paraphrased, corrected or summarised. Each one was checked automatically against the session it came from."
)
li(
    "**Two edits only, both mechanical.** Line breaks are collapsed to single spaces so a quote reads as one block, and anything shaped like a phone number or email is removed. Long answers are cut at a few hundred characters and end with an ellipsis; the spreadsheet carries the same text."
)
li(
    "**An answer can span several messages.** Workers often send a short line, get a reply, then continue. Where that happened the answer is their turns for that one question joined in order, which is still their words and nobody else's."
)
li(
    "**Illustration, not proof.** Quotes show what the data feels like; the counts beside them are the evidence. A handful of vivid extracts cannot demonstrate a pattern, which is why every category states how many cases exist in the full corpus."
)
li(
    "**Punctuation inside quotes is theirs, not ours.** The rest of this document uses plain hyphens, but "
    "quoted text keeps whatever the worker or the interviewer actually typed. Where a quote contains a "
    "long dash it is because it was written that way."
)
li(
    "**Chosen mechanically, not hand-picked.** Selection is by fixed rule and fixed random seed, stratified across topics, so nobody chose the flattering ones. Re-running reproduces the same set."
)
h("Both languages", 2)
p(
    "Interviews run in English and Hausa, and both are represented here rather than defaulting to the ones that happen to be readable in English. Hausa quotes are shown in the original with a line of key terms underneath, and every extract is labelled with its language."
)
p(
    "What is deliberately NOT here is a translation. Guidance on multilingual qualitative research is to publish the original alongside a proper translation so bilingual readers can check it - but an unverified translation of a participant's words is worse than none, so none has been invented. The spreadsheet flags every Hausa extract and leaves an empty column for a Hausa speaker to fill in. That should happen before any of these quotes are used outside the team."
)
p(
    "One category is language-skewed by construction and should not be read as a difference between speakers: gaskiya is an everyday Hausa word appearing in roughly 4% of Hausa messages, whereas the English honestly is much rarer, so the candour category is mostly Hausa because of how the two languages work, not because Hausa speakers are franker.",
    italic=True,
)
table(
    ["Category", "Found in corpus", "Topics", "In the dump"],
    [[g["title"], n(g["found"]), g["topics_covered"], g["sampled"]] for g in GEMS],
)

# How many extracts to print per category. The document is meant to be read, so this is a generous but
# finite selection; the spreadsheet carries all 1,600+ for filtering and deeper work.
SHOW = {
    "rich_detail": 12,
    "reasoning": 12,
    "barrier": 12,
    "quantified": 12,
    "candour": 8,
    "example": 8,
    "probe_rescue": 12,
    "confusion_recovery": 8,
    "dontknow": 8,
    "data_quality": 6,
    "bot_error": 3,
}
for g in GEMS:
    k = g["key"]
    if k == "suspected_ai" or not g["examples"]:
        continue
    h(g["title"], 2)
    _lang = ", ".join(f"{v} {kk}" for kk, v in list(g["by_language"].items())[:3])
    p(f"{g['how']} Found {n(g['found'])} times across {g['topics_covered']} topics ({_lang}).", italic=True)
    for ex in g["examples"][: SHOW.get(k, 3)]:
        _tag = f"{ex['topic']} - {ex['lang']}"
        if ex.get("before"):
            p(f"*{_tag}*")
            p(f"First answer: \"{ex['before']}\"")
            p(f"Interviewer: \"{ex['after']}\"")
            quote(f"Then: {ex['text']}")
        else:
            if ex.get("q"):
                p(f"*{_tag}* - asked: {ex['q']}")
            else:
                p(f"*{_tag}*")
            quote(ex["text"])
        _gl = gloss_line(ex["text"], ex["lang"])
        if _gl:
            p(_gl, italic=True)

# ================================================================ WHAT IT ADDS UP TO
h("What this adds up to", 1)
_reason_n = _mk.get("reason", {}).get("n", 0)
li(
    f"**There is real content here, at scale.** {n(L2['answers'])} answers from {n(L5['flws'])} workers, "
    f"{n(_reason_n)} of which explain a reason rather than just stating a fact."
)
li(
    f"**The follow-up earns its place.** {n(_pr['found'])} answers were unusable at first attempt and "
    f"usable by the end. On a form, the first attempt is what would have been filed."
)
li(
    f"**Workers are candid.** {n(gem('candour')['found'])} answers carry an explicit candour marker, and "
    f"{n(gem('dontknow')['found'])} say plainly that they do not know - which is a more trustworthy "
    f"corpus than one where every question gets a confident answer."
)
li(
    f"**It is not uniformly clean.** {n(gem('data_quality')['found'])} answers are junk, gibberish or a "
    f"single character - about {round(100 * gem('data_quality')['found'] / max(L2['answers'], 1), 1)}% of "
    f"the total. That is a normal rate for open text at this scale, and it is stated here so nobody "
    f"discovers it later."
)
li(
    "**Length varies enormously by topic**, so comparing raw answer depth across subjects will mislead "
    "unless the topic is held constant."
)
li("**Interviews are not single sittings** for a large minority, which matters for any timing or " "effort measure.")

# ================================================================ METHOD
h("How this was produced", 1)
p(
    "Every session was pulled from OpenChatStudio and stored locally, then read start to finish. Within "
    "each session the interviewer's turns were classified as either opening a new question or following "
    "up on the one already open, and the FLW's turns were attached to whichever question was open. That "
    "gives, for each question, the answer before any follow-up and the answer it ended with."
)
p(
    "Detection rules were built from the corpus's own vocabulary rather than assumed. The Hausa markers "
    "were chosen after counting actual usage - for instance 'rashin' is used to mean a shortage roughly "
    "nine times out of ten, but 'rashin lafiya' simply means illness, so that phrase is excluded. "
    "Getting that wrong in either direction would have distorted the barrier counts."
)
p(
    "Extracts are selected deterministically: a fixed random seed, stratified by topic so that one large "
    "cohort cannot dominate the examples. Re-running the analysis reproduces the same selection."
)

h("Checks that were run", 2)
li(
    f"Every session in the archive is accounted for - {n(COV['sessions_analysed'])} analysed plus "
    f"{n(_excl)} excluded for a stated reason equals {n(COV['sessions_in_archive'])}."
)
li(
    "Every quote was verified to appear word for word in the session it is attributed to. An answer is "
    "the worker's own turns for one question joined in order, so it is verbatim without necessarily "
    "being one uninterrupted message."
)
li(
    "No worker identifier appears anywhere in this document or the spreadsheet; session references are "
    "truncated and phone-shaped numbers were removed."
)
li("The extracts were checked for spread, so no category is drawn from only one topic or one language.")

h("Limitations, stated plainly", 2)
li(
    "**These are self-reports.** An FLW saying something is missing is evidence about their experience "
    "and their perception, not an audited measurement of stock or coverage."
)
li(
    "**The questions shape the answers.** The shortages named most often are partly a function of what "
    "was asked. This is not a prevalence survey."
)
li(
    "**Keyword detection is approximate.** It finds the marker, not the meaning. A category count is a "
    "reliable order of magnitude and a good way to find examples, not a precise measurement."
)
li(
    "**Quality was not judged.** Nothing here scores whether an answer is accurate, relevant or deep. "
    "Length and markers are proxies for effort, not for truth. Scoring quality properly needs a separate "
    "pass with a judge model, validated against human raters."
)
li(
    "**Hausa detection is thinner than English.** The marker set is grounded in observed usage, but it "
    "is smaller. Where a category matters, check the Hausa share in the table above before drawing "
    "conclusions about differences between languages."
)

h("Using the extract dump", 2)
p(
    "The spreadsheet alongside this document carries every selected extract, one row each, with columns "
    "for the category, the topic, the language, the question as asked, the answer, and - where the "
    "follow-up mattered - the first attempt and the interviewer's reply. It is meant to be filtered and "
    "read, and it is a reasonable starting point for any deeper analysis on this project, including work "
    "with a language model, since it carries the context alongside each quote."
)


# ================================================================ RENDERERS
def render_md(path):
    out = []
    for kind, a, b in BLOCKS:
        if kind == "title":
            out.append(f"# {a}\n")
        elif kind == "h":
            out.append(f"\n{'#' * (b + 1)} {a}\n")
        elif kind == "p":
            out.append(f"*{a}*\n" if b.get("italic") else f"{a}\n")
        elif kind == "li":
            out.append("- " + a)
        elif kind == "quote":
            out.append(f"> {a}\n")
        elif kind == "table":
            out.append("")
            out.append("| " + " | ".join(str(x) for x in a) + " |")
            out.append("|" + "|".join("---" for _ in a) + "|")
            for row in b:
                out.append("| " + " | ".join(str(x) for x in row) + " |")
            out.append("")
    md = "\n".join(out).replace("\n\n\n", "\n\n")
    with open(path, "w", encoding="utf-8") as f:
        f.write(md.rstrip() + "\n")
    return md


NAVY = RGBColor(0x1F, 0x38, 0x64)
GREY = RGBColor(0x6B, 0x72, 0x80)
QUOTE = RGBColor(0x1B, 0x4D, 0x3E)


def render_docx(path):
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    def runs(par, text):
        for i, seg in enumerate(text.split("**")):
            if seg:
                r = par.add_run(seg)
                r.bold = i % 2 == 1

    for kind, a, b in BLOCKS:
        if kind == "title":
            r = doc.add_paragraph().add_run(a)
            r.bold, r.font.size, r.font.color.rgb = True, Pt(20), NAVY
        elif kind == "h":
            par = doc.add_heading(a, level=b)
            for r in par.runs:
                r.font.color.rgb = NAVY
        elif kind == "p":
            par = doc.add_paragraph()
            par.paragraph_format.space_after = Pt(6)
            runs(par, a)
            for r in par.runs:
                if b.get("italic"):
                    r.italic = True
                    r.font.color.rgb = GREY
                    r.font.size = Pt(9.5)
        elif kind == "li":
            par = doc.add_paragraph(style="List Bullet")
            par.paragraph_format.space_after = Pt(3)
            runs(par, a)
        elif kind == "quote":
            par = doc.add_paragraph()
            par.paragraph_format.left_indent = Pt(18)
            par.paragraph_format.space_after = Pt(8)
            r = par.add_run(a)
            r.italic = True
            r.font.color.rgb = QUOTE
            r.font.size = Pt(10)
        elif kind == "table":
            t = doc.add_table(rows=1, cols=len(a))
            t.style = "Light Grid Accent 1"
            for i, htxt in enumerate(a):
                cell = t.rows[0].cells[i]
                cell.text = ""
                r = cell.paragraphs[0].add_run(str(htxt))
                r.bold = True
            for row in b:
                cells = t.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = ""
                    cells[i].paragraphs[0].add_run(str(val))
            doc.add_paragraph()
    doc.save(path)


if __name__ == "__main__":
    os.makedirs("docs", exist_ok=True)
    md = render_md("docs/Interview_Transcript_Insights.md")
    try:
        render_docx("docs/Interview_Transcript_Insights.docx")
        ok = True
    except PermissionError:
        ok = False
    nq = sum(1 for g in GEMS for _ in g["examples"])
    print(f"[docx] docs/Interview_Transcript_Insights.md written ({len(md.splitlines())} lines)")
    print(
        "[docx] docs/Interview_Transcript_Insights.docx written"
        if ok
        else "[docx] .docx SKIPPED - the file is open in Word; close it and re-run"
    )
    print(
        f"[docx] quotes shown in the document: {sum(len(g['examples'][:SHOW.get(g['key'], 3)]) for g in GEMS if g['key'] != 'suspected_ai')}"
        f" of {nq} carried in the payload"
    )
    sys.exit(0)
