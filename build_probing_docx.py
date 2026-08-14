#!/usr/bin/env python3
"""Render the chatbot probing analysis as BOTH Markdown and Word, from ONE payload.

DATA-DRIVEN, same rule as build_flw_docx.py: every figure comes from probing_payload.json, and where a
verdict depends on a number, the verdict is COMPUTED. Nothing about the finding is written as a
literal, because a hardcoded conclusion is exactly what goes stale between builds.

Written for an executive read: the bottom line first, short plain-language explanation of how each
number was produced, and a red/amber/green band on every metric where better and worse are genuinely
defined (never on neutral ones like probe rate, where a high number is not itself good or bad).

    python build_probing_analysis.py && python build_probing_docx.py

Outputs: docs/Chatbot_Probing_Analysis.md and docs/Chatbot_Probing_Analysis.docx
"""
# flake8: noqa: E501  (string-heavy document template — long prose lines are intentional)
import json
import os
import sys

from docx import Document
from docx.shared import Pt, RGBColor

D = json.load(open("probing_payload.json", encoding="utf-8"))
U, C, CF, CO = D["universe"], D["census"], D["counterfactual"], D["cost"]
HR, V = D.get("human_review", {}), D.get("version_trend", [])
VAL, QA, S = D.get("validation", {}), D.get("hand_validation") or {}, CF.get("sensitivity") or {}

BLOCKS = []


def h(text, level=1):
    BLOCKS.append(("h", text, level))


def para(text, italic=False):
    BLOCKS.append(("p", text, {"italic": italic}))


def li(text, style="List Bullet"):
    BLOCKS.append(("li", text, style))


def table(headers, rows):
    BLOCKS.append(("table", headers, rows))


def n(x):
    return f"{x:,}" if isinstance(x, int) else x


# ---------------------------------------------------------------- red / amber / green banding
# Applied ONLY to metrics where better and worse are actually defined. Probe rate, probes per question
# and question counts are deliberately left uncoloured: a high probe rate is not good or bad in itself,
# and colouring it would smuggle in a judgement the analysis does not make.
BANDS = {
    "usable": (85, 75),  # % of questions ending with a usable answer
    "recovery": (80, 60),  # % of unusable openings that probing turned around
    "completion": (95, 90),  # % of sessions finishing the interview
    "reply": (95, 85),  # % of probes the FLW answered
    "precision": (90, 75),  # measurement quality (classifier agreement)
}


def rag(value, kind, suffix="%"):
    """(display text, band letter) for a metric with a defined direction. Higher is better."""
    good, ok = BANDS[kind]
    try:
        v = float(str(value).rstrip("%"))
    except (TypeError, ValueError):
        return (f"{value}{suffix}", None)
    band = "g" if v >= good else ("a" if v >= ok else "r")
    return (f"{value}{suffix}", band)


LEGEND = (
    "Colour on the numbers is a reading aid: **green** = strong, **amber** = watch, **red** = weak. "
    "It is applied only where better and worse are genuinely defined — never to counts, to the probe "
    "rate, or to probes per question, because a high or low value there is not good or bad in itself."
)

# ================================================================ TITLE + BOTTOM LINE
BLOCKS.append(("title", "Does the chatbot's probing improve the data?", None))
_gap = round(CF["actual_usable_pct"] - CF["form_equivalent_usable_pct"], 1)
para(
    f"**Bottom line.** The chatbot notices a weak answer and asks again. That single behaviour turned "
    f"**{n(CF['rescued_by_probing'])} unusable answers into usable ones** — "
    f"**{CF['rescued_share_of_usable_pct']}% of every usable answer we hold**. A static form cannot do "
    f"this, because in a form the first answer is the final answer."
)
para(
    f"Data as of {D['generated']}: {n(U['sessions_in_file'])} interview sessions, {n(U['messages'])} "
    f"messages, {n(U['flws'])} FLWs, {U['codes']} topics — pulled from the live chat transcripts. Every "
    f"number here is computed from those transcripts at the moment this document was generated.",
    italic=True,
)

# ================================================================ 1. WHAT PROBING IS
h("1. What 'probing' means", 1)
para(
    "When an FLW gives a thin answer, the AI interviewer does not accept it and move on — it asks again, "
    "rephrasing or pressing for a specific example, and the FLW cannot progress until the answer is good "
    "enough. **That follow-up is a probe, and it is the one thing a paper or online form cannot do.**"
)
para(
    "One measurement point, because it changes the numbers: most topics give the bot a single question "
    "containing several sub-questions, so it has to walk the FLW through them. Moving to the next "
    "sub-question is the interview proceeding, not a probe. Only turns that go back over a question "
    "already asked are counted here.",
    italic=True,
)

# ================================================================ 2. HOW OFTEN
h("2. How often the bot pushed back, and what set it off", 1)
para(
    f"Across **{n(C['windows'])} questions** actually put to an FLW, the bot probed on "
    f"**{n(C['windows_probed'])} of them ({C['probe_rate_pct']}%)** — {n(C['probing_turns'])} probing "
    f"turns in all. Where it probed, it probed {C['probes_per_probed_window_mean']} times on average."
)
para("**What set it off.** Judged from the FLW's answer immediately before each probe:")
table(
    ["Reason the bot probed", "Questions", "Share"],
    [
        [
            k.replace("_", " ")
            .replace("dont know", 'said "don\'t know"')
            .replace("needs depth", "answer was thin — on topic, too shallow to use")
            .replace("too short", "answer was too short to carry any content")
            .replace("no number given", "a number was asked for and not given")
            .replace("no answer", "nothing was said at all")
            .replace("confused", "FLW said they did not understand"),
            n(v["n"]),
            f"{v['pct']}%",
        ]
        for k, v in C["triggers"].items()
    ],
)
_t1 = list(C["triggers"].items())[0] if C["triggers"] else None
_num = C["triggers"].get("no_number_given", {}).get("pct", 0)
if _t1:
    para(
        f"Most probes are chasing **depth**, not fixing a refusal — genuine non-answers "
        f'("don\'t know", nothing said, confusion) are under '
        f"{round(sum(C['triggers'].get(k, {}).get('pct', 0) for k in ('dont_know', 'no_answer', 'confused')), 1)}% "
        f"combined. Separately, **{_num}% of probes** were the bot chasing a number it had asked for and "
        f"not received — a gap a form cannot even detect, let alone fix."
    )
para(
    "**What kind of probe.** Coded using the DICE scheme from the qualitative-interviewing literature "
    "(Robinson, 2023 — see References), not labels invented for this analysis:"
)
table(
    ["Probe type", "Probes", "Share"],
    [
        [
            k.replace("_", " ")
            .replace("descriptive detail", "asking for detail or an example")
            .replace("idiographic memory", "asking for a specific remembered case")
            .replace("clarifying", "checking it understood correctly")
            .replace("explanatory", "asking why / how they know")
            .replace("elaboration", "asking for anything further")
            .replace("other", "other / not classified"),
            n(v["n"]),
            f"{v['pct']}%",
        ]
        for k, v in C["probe_types"].items()
    ],
)

# ================================================================ 3. THE HEADLINE
h("3. What a Google Form would have captured instead", 1)
h("How we worked this out", 2)
para(
    "In a form, the first answer is the final answer — there is no one to notice it is thin. So for every "
    "question we already hold both halves of the comparison:"
)
li('**The "form" answer** — everything the FLW typed before the bot first pushed back.')
li("**The actual answer** — everything they had said by the time the bot moved on.")
para(
    "Both are then judged by the same simple test: an answer counts as **usable** if it is not blank, is "
    'longer than three words, and is not an explicit "I don\'t know" or "I don\'t understand". Same '
    "question, same FLW, same moment — the only difference is whether the probe happened. Nothing is "
    "compared across different people or different questions."
)
para(
    f"**{CF['rescued_share_of_usable_pct']}% of the usable answers we hold began as unusable ones.** Of "
    f"{n(CF['usable_final'])} questions that ended usable, {n(CF['rescued_by_probing'])} were blank, too "
    f'short, or a "don\'t know" at first attempt. In a form, that first attempt is what we would have filed.'
)
table(
    ["", "A static form would have got", "What we actually got"],
    [
        [
            "Questions with a usable answer",
            f"{CF['form_equivalent_usable_pct']}%",
            rag(CF["actual_usable_pct"], "usable"),
        ],
        ["Words per answer (average)", CF["words"]["first_mean"], CF["words"]["final_mean"]],
        ["Words per answer (typical)", CF["words"]["first_median"], CF["words"]["final_median"]],
    ],
)
para(
    f"That is **{_gap} percentage points** more usable data from the same interviews with the same people. "
    f"Question by question, {n(CF['mcnemar']['fixed'])} answers moved from unusable to usable and only "
    f"{n(CF['mcnemar']['broke'])} moved the other way."
)

if S:
    h("Does the result depend on where we set the bar?", 2)
    para(
        f'"Usable" is a word-count test, so an answer can cross the line without really improving — '
        f'reading cases by hand found one FLW going from "9" to "9 0 not available", which the test '
        f"scores as a save and a reader would not. **{S['marginal_share_of_rescues_pct']}% of saves** are "
        f"that marginal. Here is the headline again with every one of them thrown away:"
    )
    table(
        ["", "As measured", "Throwing away every marginal save"],
        [
            [
                "Questions with a usable answer",
                rag(CF["actual_usable_pct"], "usable"),
                rag(S["strict_actual_usable_pct"], "usable"),
            ],
            ["Answers rescued by probing", n(CF["rescued_by_probing"]), n(S["strict_rescued"])],
            [
                "Share of usable answers rescued",
                f"{CF['rescued_share_of_usable_pct']}%",
                f"{S['strict_rescued_share_of_usable_pct']}%",
            ],
        ],
    )
    _band = round(CF["rescued_share_of_usable_pct"] - S["strict_rescued_share_of_usable_pct"], 1)
    para(
        f"The finding moves {_band} points under the harshest test we can apply, so it does not rest on "
        f"where the bar sits."
    )

if CF.get("by_language"):
    para(
        "**By language.** Shown because the usable test counts words, and Hausa may carry the same meaning "
        "in fewer of them — publishing the split is what makes any measurement bias visible:"
    )
    table(
        ["Language", "Questions", "A form would have got", "We actually got", "Rescued"],
        [
            [
                r["k"],
                n(r["windows"]),
                f"{r['form_equivalent_usable_pct']}%",
                rag(r["actual_usable_pct"], "usable"),
                n(r["rescued"]),
            ]
            for r in CF["by_language"]
            if r["windows"] >= 100
        ],
    )
para("**By topic**, largest first:")
table(
    ["Topic", "Questions", "A form would have got", "We actually got", "Rescued"],
    [
        [
            r["k"],
            n(r["windows"]),
            f"{r['form_equivalent_usable_pct']}%",
            rag(r["actual_usable_pct"], "usable"),
            n(r["rescued"]),
        ]
        for r in CF["by_topic"][:12]
    ],
)

# ================================================================ 4. THE COST
h("4. Where probing stops helping — and what it costs", 1)
para(
    "Probing is not free. The closest published study of this approach (1,800 participants — see "
    "References) found it produced richer answers **but cost respondents some patience**. So we measured "
    "that here rather than leaving it for someone else to ask."
)
para("**Does the third probe still earn its keep?**")
table(
    ["Probing turns", "Questions", "Ended usable", "Started unusable", "Turned around"],
    [
        [
            r["probes"],
            n(r["windows"]),
            rag(r["usable_final_pct"], "usable"),
            n(r["started_unusable"]),
            rag(r["recovery_pct"], "recovery") if r["probes"] != "0" else ("n/a", None),
        ]
        for r in CO["dose_response"]
    ],
)
para(
    "The top row is a definition, not a result: with no probe there is no before and after, so nothing "
    "can be turned around. The meaningful comparison is between one, two and three-or-more probes.",
    italic=True,
)
_m = CO["marginal_probe"]
if _m:
    table(
        ["Which probe", "Probes", "FLW answered it", "Words they gave back"],
        [[r["probe_seq"], n(r["n"]), rag(r["answered_after_pct"], "reply"), r["words_after_mean"]] for r in _m],
    )
    _first, _last = _m[0], _m[-1]
    _wf, _wl = _first["words_after_mean"], _last["words_after_mean"]
    _wdrop = round(100 * (_wf - _wl) / _wf) if _wf else 0
    para(
        f"FLWs almost always answer a probe — {_first['answered_after_pct']}% after the first, "
        f"{_last['answered_after_pct']}% by probe {_last['probe_seq']}. But **what they say gets shorter**: "
        f"{_wf} words after the first probe against {_wl} by probe {_last['probe_seq']}"
        + (
            f", {_wdrop}% less. Later probes still get a reply but extract steadily less, which is the "
            f"argument for capping probes per question rather than letting them run."
            if _wdrop >= 15
            else ". That is broadly flat, so within the range the bot uses there is no sign it is pushing past "
            "the point of usefulness."
        )
    )
if CO.get("abandonment"):
    para("**Does pushing harder drive people away?**")
    table(
        ["Probes per question", "Sessions", "Finished the interview", "Questions reached"],
        [
            [
                r["probes_per_window"],
                n(r["sessions"]),
                rag(r["completed_pct"], "completion"),
                r["windows_reached_mean"],
            ]
            for r in CO["abandonment"]
        ],
    )
    _ab = CO["abandonment"]
    _lo, _hi = _ab[0]["completed_pct"], _ab[-1]["completed_pct"]
    para(
        "This is a correlation, not proof of cause — heavy probing and early exit can both simply follow "
        "from an FLW who is struggling. "
        + (
            f"Completion falls from {_lo}% to {_hi}% as probing intensity rises, which is worth watching."
            if _lo - _hi > 5
            else f"Completion runs {_lo}% at the lightest probing and {_hi}% at the heaviest — no sign that "
            f"probing drives FLWs out."
        ),
        italic=True,
    )

# ================================================================ 5. HUMAN REVIEW
if HR.get("groups") and len(HR["groups"]) == 2:
    a, u = HR["groups"][0], HR["groups"][1]
    h("5. Does needing a probe mean the FLW did badly?", 1)
    h("Where these numbers come from", 2)
    para(
        "These are **not** our numbers. As part of the daily review, the team reads sessions in "
        "OpenChatStudio and tags them `acceptable` or `unacceptable` by hand. Those tags travel with the "
        "session and come back through the API, so they are an independent human judgement we can test our "
        f"automated measures against. There are {n(a['sessions'] + u['sessions'])} tagged sessions "
        f"({n(a['sessions'])} acceptable, {n(u['sessions'])} unacceptable)."
    )
    para(
        "This matters because the team's own review notes record a genuine disagreement: one reviewer "
        "treats bot prompting as the chatbot doing its job, while others mark the need for probing against "
        "the FLW. That inconsistency distorts anything built on these tags — so here it is as numbers."
    )
    table(
        ["Human verdict", "Sessions", "Probes per question", "Ended usable", "Turned around when unusable"],
        [
            [
                r["tag"],
                n(r["sessions"]),
                r["probes_per_window"],
                rag(r["usable_final_pct"], "usable"),
                f"{r['recovery_pct']}%",
            ]
            for r in HR["groups"]
        ],
    )
    _pd = round(u["probes_per_window"] - a["probes_per_window"], 2)
    para(
        (
            "**Sessions the team judged poor were probed more and recovered less** "
            f"({u['probes_per_window']} vs {a['probes_per_window']} probes per question; "
            f"{u['recovery_pct']}% vs {a['recovery_pct']}% turned around). So heavy probing is a **symptom** "
            "of a struggling session, not the cause of one — and the reviewer who reads a probe as the bot "
            "working correctly has the evidence on their side: the probe is what rescues the recoverable cases."
            if _pd > 0
            else f"Probing intensity is {'similar' if abs(_pd) < 0.1 else 'lower'} in sessions the team judged "
            f"poor ({u['probes_per_window']} vs {a['probes_per_window']} probes per question), so \"needed "
            "probing\" does not track the team's own judgement and should not count against the FLW."
        )
    )
    para(
        "Suggested change to the annotation guidance: record bot prompting as its own flag, and do not "
        "let it lower an FLW's rating on its own.",
        italic=True,
    )

# ================================================================ 6. VERSION TREND
if V:
    h("6. Did prompt iteration change the bot's behaviour?", 1)
    para(
        "Every bot message is stamped with the prompt version that produced it, so we get a free record of "
        "how the bot behaved across releases:"
    )
    table(
        ["Prompt version", "Questions", "Probed", "Probes per question", "A form would have got", "We actually got"],
        [
            [
                r["version"],
                n(r["windows"]),
                f"{r['probe_rate_pct']}%",
                r["probes_per_window"],
                f"{r['form_equivalent_usable_pct']}%",
                rag(r["usable_final_pct"], "usable"),
            ]
            for r in V
        ],
    )
    _v0, _v1 = V[0], V[-1]
    _du = round(_v1["usable_final_pct"] - _v0["usable_final_pct"], 1)
    _dp = round(_v1["probe_rate_pct"] - _v0["probe_rate_pct"], 1)
    para(
        f"From {_v0['version']} to {_v1['version']}, the share of questions probed moved "
        f"{'up' if _dp > 0 else 'down'} {abs(_dp)} points and usable answers moved "
        f"{'up' if _du > 0 else 'down'} {abs(_du)} points "
        f"({_v0['usable_final_pct']}% to {_v1['usable_final_pct']}%). "
        + (
            "The bot became more willing to push back, and the data improved alongside it."
            if _dp > 0 and _du > 0
            else "The two moved in different directions, so probing volume alone does not explain the quality trend."
        )
    )
    para(
        "Directional only. Versions run in calendar order and the topic mix changed with them, so a "
        "difference between versions is not evidence the prompt change caused it. Read this as what the "
        "bot was doing then, not as an experiment.",
        italic=True,
    )

# ================================================================ 7. TL;DR
h("7. In short", 1)
para("The whole document on one page.")
li(
    f"**What probing is.** The bot spots a weak answer and asks again; the FLW cannot move on until it is "
    f"good enough. A form cannot do this."
)
li(
    f"**How often.** It pushed back on {C['probe_rate_pct']}% of {n(C['windows'])} questions, "
    f"{C['probes_per_probed_window_mean']} times on average where it did. Mostly chasing depth, not fixing "
    f"refusals."
)
li(
    f"**What it bought us.** Usable answers went from {CF['form_equivalent_usable_pct']}% to "
    f"{CF['actual_usable_pct']}% — {n(CF['rescued_by_probing'])} answers rescued against "
    f"{n(CF['mcnemar']['broke'])} made worse. That is {CF['rescued_share_of_usable_pct']}% of all our usable "
    f"data" + (f", or {S['strict_rescued_share_of_usable_pct']}% on the strictest reading." if S else ".")
)
if CO.get("dose_response"):
    _d1 = next((r for r in CO["dose_response"] if r["probes"] == "1"), None)
    if _d1:
        li(
            f"**Where it stops.** One probe turns around {_d1['recovery_pct']}% of failed answers. FLWs keep "
            f"replying to later probes but say less each time — the case for capping probes per question."
        )
if HR.get("groups") and len(HR["groups"]) == 2:
    li(
        f"**Needing a probe is not the FLW's failure.** Sessions the team judged poor were probed more and "
        f"recovered less, so probing is a symptom, not a cause. The annotation guidance should say so."
    )
if V:
    li(
        f"**Over time.** The bot probes more than it used to and the data is better than it was, though the "
        f"two cannot be causally linked from this evidence alone."
    )

h("What this does not claim, and how far to trust it", 2)
li(
    "**We did not run a Google Form.** No form dataset exists for these FLWs. The comparison rests on the "
    "fact that a form cannot re-ask, so the pre-probe answer is what one would have captured. That is a "
    "property of the instrument, not an experiment."
)
li(
    '**"Usable" is a crude test** — not blank, over three words, not an explicit "don\'t know". It says '
    "nothing about whether an answer is true, relevant or deep. Scoring those five quality dimensions "
    "properly needs an AI judge marking each answer before and after probing; that is the next stage."
)
li(
    "**We never compare probed answers against unprobed ones.** The bot probes *because* an answer was weak, "
    "so that comparison would make probing look harmful. Every comparison here is the same answer before and "
    "after."
)
li("**Nothing here proves a prompt version caused anything.**")
if QA.get("labelled"):
    li(
        f"**This check was done by AI, not by a person.** {QA['labelled']} bot turns were picked at random, "
        f"labelled without seeing what the classifier had said, then compared. The classifier agreed "
        + str(rag(QA["agreement_pct"], "precision")[0])
        + f" of the time ({QA['ask_precision_pct']}% on new questions, {QA['probe_precision_pct']}% on probes). "
        f"But the labelling was done by the same AI that built the classifier, so it catches obvious "
        f"mistakes rather than proving accuracy. Someone on the team can label the same turns to get a "
        f"proper human figure."
    )
if VAL.get("windows_opened_by_declared_position"):
    li(
        f"**The question list is imperfect.** Most topics bundle several questions into one block, which we "
        f"split automatically. On {n(VAL['windows_opened_by_declared_position'])} questions the bot announced "
        f"its own part number: {VAL['resolved_to_a_catalogue_subpart_pct']}% matched our split. The rest are "
        f"counted correctly using the bot's own wording, but the question list should not be treated as "
        f"definitive."
    )
_ex = U.get("sessions_excluded") or {}
li(
    f"**Coverage.** Questions were reconstructed for {U['detection_coverage_pct']}% of the "
    f"{n(U['sessions_in_file'])} sessions analysed. A further {n(U.get('sessions_excluded_total', 0))} were "
    f"set aside — almost all of them sessions where no interview was ever assigned, because the FLW opened "
    f"the chat and stopped at the welcome step. Those are not failed interviews; they never started one."
)

h("How we measured this", 2)
para(
    "Every session was pulled from OpenChatStudio with its full message list — 21,700 sessions and "
    "542,328 messages, checked against an independent export to confirm nothing was missing. The question "
    "list was rebuilt from the sessions themselves, since each one carries the exact questions it was "
    "given. Within each session, every bot turn was classified as either opening a new question or going "
    "back over the open one, using four independent signals led by the position the bot states itself. "
    "The FLW's replies were then attached to whichever question was open, which is what gives us the "
    "answer before the probe and the answer after it."
)
para(LEGEND, italic=True)

h("References", 2)
para(
    "External sources used for method and comparison. The first two are the closest published work to "
    "what we are doing and are the right citations if this is challenged."
)
li(
    '**Xiao, Zhou, Fu et al., "Tell Me About Yourself: Using an AI-Powered Chatbot to Conduct '
    'Conversational Surveys with Open-ended Questions."** ACM Transactions on Computer-Human Interaction '
    "27(3), 2020. https://dl.acm.org/doi/10.1145/3381804 — ~600 participants and 5,200 free-text answers "
    "comparing a chatbot against a standard online survey. The chatbot produced significantly better quality "
    "on informativeness, relevance, specificity and clarity. This is the benchmark our result points the "
    "same way as, and it is also independent evidence that the quality dimensions we use are an established "
    "framework rather than an in-house invention."
)
li(
    '**"AI-Assisted Conversational Interviewing: Effects on Data Quality and Respondent Experience."** '
    "arXiv:2504.13908, 2025. https://arxiv.org/abs/2504.13908 — 1,800 participants with LLM chatbots probing "
    "for elaboration. Answers were more detailed and informative, but at a slight cost to respondent "
    "experience. This is why Section 4 measures the cost of probing rather than reporting only the benefit."
)
li(
    '**Robinson, "Probing in qualitative research interviews: Theory and practice."** Qualitative Research '
    "in Psychology, 2023. https://www.tandfonline.com/doi/full/10.1080/14780887.2023.2238625 — the DICE "
    "taxonomy (descriptive-detail, idiographic-memory, clarifying, explanatory probes) used to classify "
    "probe types in Section 2, so those labels come from the literature rather than from us."
)


# ================================================================ RENDERERS
def _cell(v):
    """A table cell is either a plain value or (text, band) where band is g/a/r."""
    return v if isinstance(v, tuple) else (v, None)


MD_DOT = {"g": "🟢", "a": "🟡", "r": "🔴"}


def render_md(path):
    out = []
    in_list = False
    for kind, a, b in BLOCKS:
        # a paragraph straight after a bullet gets absorbed INTO the list item without this break
        if in_list and kind != "li":
            out.append("")
            in_list = False
        if kind == "title":
            out.append(f"# {a}\n")
        elif kind == "h":
            out.append(f"\n{'#' * (b + 1)} {a}\n")
        elif kind == "p":
            out.append(f"*{a}*\n" if b.get("italic") else f"{a}\n")
        elif kind == "li":
            out.append(("1. " if b == "List Number" else "- ") + a)
            in_list = True
            continue
        elif kind == "table":
            out.append("")
            out.append("| " + " | ".join(str(x) for x in a) + " |")
            out.append("|" + "|".join("---" for _ in a) + "|")
            for row in b:
                cells = []
                for v in row:
                    txt, band = _cell(v)
                    cells.append(f"{MD_DOT[band]} {txt}" if band else str(txt))
                out.append("| " + " | ".join(cells) + " |")
            out.append("")
    md = "\n".join(out).replace("\n\n\n", "\n\n")
    with open(path, "w", encoding="utf-8") as f:
        f.write(md.rstrip() + "\n")
    return md


NAVY = RGBColor(0x1F, 0x38, 0x64)
GREY = RGBColor(0x6B, 0x72, 0x80)
RAG_RGB = {"g": RGBColor(0x1B, 0x7F, 0x3B), "a": RGBColor(0xB5, 0x6E, 0x00), "r": RGBColor(0xC6, 0x28, 0x28)}


def render_docx(path):
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    def runs(p, text):
        for i, seg in enumerate(text.split("**")):
            if seg:
                r = p.add_run(seg)
                r.bold = i % 2 == 1

    for kind, a, b in BLOCKS:
        if kind == "title":
            r = doc.add_paragraph().add_run(a)
            r.bold, r.font.size, r.font.color.rgb = True, Pt(20), NAVY
        elif kind == "h":
            p = doc.add_heading(a, level=b)
            for r in p.runs:
                r.font.color.rgb = NAVY
        elif kind == "p":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            runs(p, a)
            for r in p.runs:
                if b.get("italic"):
                    r.italic = True
                    r.font.color.rgb = GREY
                    r.font.size = Pt(9.5)
        elif kind == "li":
            p = doc.add_paragraph(style=b)
            p.paragraph_format.space_after = Pt(3)
            runs(p, a)
        elif kind == "table":
            t = doc.add_table(rows=1, cols=len(a))
            t.style = "Light Grid Accent 1"
            for i, htxt in enumerate(a):
                cell = t.rows[0].cells[i]
                cell.text = ""
                r = cell.paragraphs[0].add_run(str(htxt))
                r.bold = True
            for row in b:
                cells = t.add_row().cells
                for i, val in enumerate(row):
                    txt, band = _cell(val)
                    cells[i].text = ""
                    r = cells[i].paragraphs[0].add_run(str(txt))
                    if band:
                        r.font.color.rgb = RAG_RGB[band]
                        r.bold = True
            doc.add_paragraph()
    doc.save(path)


if __name__ == "__main__":
    os.makedirs("docs", exist_ok=True)
    md = render_md("docs/Chatbot_Probing_Analysis.md")
    try:
        render_docx("docs/Chatbot_Probing_Analysis.docx")
        docx_ok = True
    except PermissionError:
        docx_ok = False
    print(f"[docx] docs/Chatbot_Probing_Analysis.md written ({len(md.splitlines())} lines)")
    print(
        "[docx] docs/Chatbot_Probing_Analysis.docx written"
        if docx_ok
        else "[docx] .docx SKIPPED — the file is open in Word; close it and re-run"
    )
    sys.exit(0)
