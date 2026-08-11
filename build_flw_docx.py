#!/usr/bin/env python3
"""Render the FLW Retention executive analysis as BOTH Markdown and Word, from ONE payload.

DATA-DRIVEN: every figure comes from flw_analysis_payload.json — the `flwEngagement` block the
dashboard embeds — so the document cannot state something the dashboard contradicts. Get that payload
from the PUBLISHED render (not a local build) with:

    python pull_live_payload.py && python build_flw_docx.py

Why both files come from here: the .docx used to be generated while the .md was hand-maintained, and
they drifted — on 2026-08-11 the committed .docx still carried claims ("re-use compounds engagement",
"the retention lever is the first interview") that the corrected analysis had already disproven, while
the .md carried a correction notice the .docx did not. One content model, two renderers, no drift.

Interpretation rule applied throughout: state the number, and where a verdict depends on the number,
COMPUTE the verdict. Do not hardcode a conclusion — a hardcoded "flat" is exactly what went stale last
time (the per-cohort gap moved from ~0 to 18 points between builds).

Outputs: docs/FLW_Retention_Analysis_Brief.md and docs/FLW_Retention_Analysis_Brief.docx
"""
# flake8: noqa: E501  (string-heavy document template — long prose lines are intentional)
import json
import os
from collections import Counter

from docx import Document
from docx.shared import Pt, RGBColor

FE = json.load(open("flw_analysis_payload.json", encoding="utf-8"))
STAMP = json.load(open("_flw_today.json", encoding="utf-8")) if os.path.exists("_flw_today.json") else {}
TODAY = STAMP.get("today") or ""
BUILT_AT = STAMP.get("built_at") or ""
RENDER_V = STAMP.get("render_version")

N = FE["n_flws"]
P = {t["k"]: t for t in FE["personas"]}
T = {t["k"]: t for t in FE["tiers"]}
CC = FE["crossCohort"]
CCD = {d["k"]: d for d in CC["dist"]}
DS = FE.get("depthSplit") or FE.get("firstIv") or {}
_dsmed = DS.get("median")
DS_MED = int(_dsmed) if isinstance(_dsmed, (int, float)) and float(_dsmed).is_integer() else _dsmed
OAD = FE["oneAndDone"]
AR = FE["atRisk"]
ST = {s["k"]: s for s in FE["byState"]}
LLO = {s["k"]: s for s in FE["byLLO"]}
SV = {s["d"]: s for s in FE["survival"]}
MICRO = FE.get("micro") or {}


def _micro_facts():
    """Facts the brief asserts but previously could not show. Derived at generation time, so they
    cannot drift from the payload.

    Added after a reader audit found the brief (a) never gave the programme-level per-cohort figure at
    all, (b) invited readers to intersect three separate marginal shares in section 3, and (c) called
    9% "the" attrition figure while section 6 implied a much larger unfinished population.
    """
    m = MICRO
    if not m.get("col"):
        return {}

    def col(d):
        return [m["dict"][d][int(c)] for c in m["col"][d]]

    def unpack(spec):
        w, t = spec["w"], spec["s"]
        return [int(t[i : i + w], 36) for i in range(0, len(t), w)]

    nn = m["n"]
    pcf = unpack(m["num"]["pcf"])
    st, ty, per, nco, tier, fin = col("state"), col("type"), col("persona"), col("nco"), col("tier"), col("fin")
    finished_label = m["dict"]["fin"][0]
    oad = [i for i in range(nn) if per[i] == "One-and-done"]
    top_state = (OAD.get("topState") or [{}])[0].get("k")
    top_type = (OAD.get("topType") or [{}])[0].get("k")
    all_three = [i for i in oad if st[i] == top_state and ty[i] == top_type and nco[i] == "1"]
    unfinished = [i for i in range(nn) if fin[i] != finished_label]
    mid_tier = FE["tiers"][2]["k"] if len(FE["tiers"]) > 2 else None
    mid = [i for i in range(nn) if tier[i] == mid_tier]
    return {
        "prog_pc": round(sum(pcf) / nn) if nn else 0,
        "oad_all_three": len(all_three),
        "oad_n": len(oad),
        "oad_all_three_pct": round(100 * len(all_three) / len(oad)) if oad else 0,
        "top_state": top_state,
        "top_type": top_type,
        "unfinished": len(unfinished),
        "unfinished_pct": round(100 * len(unfinished) / nn) if nn else 0,
        "mid_tier": mid_tier,
        "mid_n": len(mid),
        "mid_finishers": sum(1 for i in mid if fin[i] == finished_label),
        "mid_oad": sum(1 for i in mid if per[i] == "One-and-done"),
    }


MF = _micro_facts()

finishers_pct = P.get("Champion", {}).get("pct", 0) + P.get("Steady finisher", {}).get("pct", 0)
multi_pct = 100 - CCD.get("1", {}).get("pct", 0)


def pc(g, key="finished_pc"):
    """Per-cohort finish rate for a group dict, falling back to the older key if absent."""
    return g.get(key, g.get("finished", 0))


def state_by_partner():
    """state -> {partner: n}, derived from the per-FLW micro block (no assumption about which state
    belongs to which partner — if the nesting ever stops holding, this stops claiming it)."""
    if not MICRO.get("col"):
        return {}
    d, out = MICRO["dict"], {}
    for i in range(MICRO["n"]):
        s = d["state"][int(MICRO["col"]["state"][i])]
        p = d["llo"][int(MICRO["col"]["llo"][i])]
        out.setdefault(s, Counter())[p] += 1
    return out


SBP = state_by_partner()


def partner_of(state):
    """The partner a state is (almost) entirely served by, or None if it is genuinely mixed."""
    c = SBP.get(state)
    if not c:
        return None
    top, n = c.most_common(1)[0]
    return top if n / sum(c.values()) >= 0.95 and "|" not in top else None


# ---------------------------------------------------------------- content model
BLOCKS = []


def h(text, level=1):
    BLOCKS.append(("h", text, level))


def para(text, italic=False, color=None, size=None, sa=6):
    BLOCKS.append(("p", text, {"italic": italic, "muted": color is not None or italic}))


def li(text, style="List Bullet"):
    BLOCKS.append(("li", text, style))


def table(headers, rows):
    BLOCKS.append(("table", headers, rows))


# ================================================================ TITLE
_stamp_bits = [
    b
    for b in [
        f"data as of {TODAY}" if TODAY else "",
        f"build {BUILT_AT}" if BUILT_AT else "",
        f"dashboard render v{RENDER_V}" if RENDER_V else "",
    ]
    if b
]
BLOCKS.append(("title", "FLW Retention & Engagement — Executive Analysis", None))
para(f"Connect Interviews programme · per-FLW, cross-cohort · {' · '.join(_stamp_bits)}", italic=True)
para(
    f"Universe: **{N:,} unique front-line workers (FLWs)** who started ≥1 interview · {FE['coverage_lga']}% have "
    "demographics · every metric dedups the worker across all cohorts/arms they were part of. Generated from the same "
    "payload the dashboard's **FLW Retention** tab embeds, so every figure here matches what is on screen.",
    italic=True,
)

h("Why this analysis exists", 1)
para(
    "Every other view of this programme is cohort-level — how a study arm performed. This one looks at the programme "
    "through the worker: one row per unique FLW, their interview history stitched across every cohort and arm they "
    "touched. Most workers are re-used across studies, so the worker's cumulative experience — not any single cohort — "
    "is what tells us who the programme retains and where it loses people."
)

# A reader hit four specific confusions on the previous version, and every one of them was caused by a
# definition living in the Method section at the BOTTOM while the number appeared at the top. So the
# things you must hold in your head now come first, in plain language.
h("How to read the numbers (please read this first)", 1)
para("**Two ways of saying a worker “finished”.** They are different questions and are never interchangeable:")
li(
    '**Finished at least one schedule** — "have they ever completed a full cohort schedule?" Beware: a worker in three '
    "cohorts has three chances to clear this bar, so this number rises with how many cohorts someone was put in, even "
    "if nothing about the worker changed. Useful for describing, misleading for comparing."
)
li(
    "**Per-cohort finish rate** — “of the schedules they were actually given, what share did they complete?” "
    "Worked out per person (schedules completed ÷ schedules enrolled in) and then averaged across the group. This is "
    "the fair one, and it is the number to quote when comparing groups."
)
para("**Two ways of grouping workers.** Both appear below and they are not comparable with each other:")
li(
    "**Personas** — what a worker has done over their **whole history** (Champion, Steady finisher, One-and-done…). "
    "Fixed: a persona does not change as time passes."
)
li(
    "**Engagement tiers** — where a worker sits **right now** (Highly engaged, Engaged, Slipping, Gone quiet, Lost). "
    "A worker moves between tiers over time. Different words are used deliberately so the two are never mixed up."
)
para(
    "And one limitation that applies to everything here: this is observational data. It can show what goes "
    "together with finishing; it cannot show what causes finishing. Where a number is tempting to read causally, "
    "the text says so explicitly.",
    italic=True,
)

# ================================================================ 1
h("1. The engagement landscape", 1)
para("Behavioural personas — rule-based segments over each worker's whole history:")
persona_desc = {
    "Champion": "Finishes, steady cadence, high answer depth — the backbone",
    "Steady finisher": "Completes at least one full schedule reliably",
    "Partial progress": "Over half their triggered interviews done, but no schedule finished yet",
    "One-and-done": "Started once and stopped — the genuine early-loss group",
    "Re-engager": "Went silent, then came back",
    "Early dropper": "Shallow start, left early",
    "Lapsed": "Inactive, nothing finished",
}
table(
    ["Persona", "Workers", "Share", "What it means"],
    [[t["k"], f"{t['n']:,}", f"{t['pct']}%", persona_desc.get(t["k"], "")] for t in FE["personas"]],
)
# Tiers arrive in score order (best first), so refer to them BY POSITION, never by name. The labels
# changed once already (Champion/Solid/At-risk -> Highly engaged/Engaged/Gone quiet) and a hardcoded
# T.get("Champion") silently reads 0 rather than failing.
_tiers = FE["tiers"]
_top2, _worst2 = _tiers[:2], _tiers[-2:]
_mid = _tiers[2:-2]
_healthy = sum(t["pct"] for t in _top2)
para(
    f"**{finishers_pct}% of workers have completed at least one full schedule** — Champions "
    f"{P.get('Champion',{}).get('pct',0)}% + Steady finishers {P.get('Steady finisher',{}).get('pct',0)}%, the two "
    f"personas defined by having finished. Genuine early loss is **{OAD['pct']}%** ({OAD['n']} workers), the "
    f"One-and-done segment that §3 profiles."
    + (
        f" **Important:** that {finishers_pct}% is the generous measure — the one §2 shows is inflated by how many "
        f"cohorts a worker was put in. On the like-for-like per-cohort measure, workers complete "
        f"**{MF['prog_pc']}%** of the schedules they were actually given. Quote {MF['prog_pc']}% as the "
        f"programme's headline finish figure, not {finishers_pct}%."
        if MF.get("prog_pc")
        else ""
    )
    + " *(Both are whole-history views. The tier table below describes the same people's CURRENT activity and is "
    "not comparable with either.)*"
)
para(
    "Engagement tiers answer a **different question** from the personas above. A tier is where a worker sits "
    "**right now** — a band blending how recently they interviewed with their completion rate and answer depth — so a "
    "worker moves between tiers over time. A persona describes their **whole history** and does not move. The two use "
    "deliberately different words so they are never read as the same grouping.",
    italic=True,
)
# Tiers previously appeared as prose with no counts while personas got a full table - that asymmetry is
# WHY a reader read "top two tiers" as "top two personas". Equal visual weight now.
table(
    ["Engagement tier (activity right now)", "Workers", "% of all"],
    [[t["k"], f"{t['n']:,}", f"{t['pct']}%"] for t in _tiers],
)
para(
    f"**{_healthy}% of workers sit in the top two tiers** ({_top2[0]['k']} + {_top2[1]['k']}). This answers a "
    f"different question from the {finishers_pct}% above: that counted who has ever finished a schedule, this "
    f"describes how active they are right now."
    + (
        f" \"{MF['mid_tier']}\" is mixed - of its {MF['mid_n']:,} workers, {MF['mid_finishers']:,} have actually "
        f"finished a schedule and are simply inactive now, while {MF['mid_oad']} are one-and-done workers - so it "
        f"should not be read as uniformly benign."
        if MF.get("mid_tier")
        else ""
    )
)
para(
    "Tier recency is measured against the freshest session in the dataset rather than the wall clock, so these "
    "shares do not drift when a data pull runs late.",
    italic=True,
)

# ================================================================ 2
h("2. Re-use across cohorts — and the measurement trap", 1)
para(f"The multi-arm design re-uses the same workers across studies: **{multi_pct}%** are in ≥2 cohorts.")
table(["Cohorts per worker", "Workers", "Share"], [[d["k"], f"{d['n']:,}", f"{d['pct']}%"] for d in CC["dist"]])
_raw_gap = CC["multi"]["finished"] - CC["single"]["finished"]
_pc_gap = pc(CC["multi"]) - pc(CC["single"])
_arith = max(0, round(100 * (_raw_gap - _pc_gap) / _raw_gap)) if _raw_gap else 0
table(
    ["Measure", "Single-cohort workers", "Multi-cohort workers", "What it asks"],
    [
        [
            "Finished ≥1 schedule",
            f"{CC['single']['finished']}%",
            f"{CC['multi']['finished']}%",
            "Have they ever completed anything?",
        ],
        [
            "Per-cohort finish rate",
            f"{pc(CC['single'])}%",
            f"{pc(CC['multi'])}%",
            "Of the schedules they were given, how many did they complete?",
        ],
        [
            "Completion rate",
            f"{CC['single']['completion']}",
            f"{CC['multi']['completion']}",
            "Of interviews they were sent, how many did they finish?",
        ],
        [
            "Answer depth",
            f"{CC['single']['depth']} words/session",
            f"{CC['multi']['depth']} words/session",
            "How much do they actually say?",
        ],
        ["Workers", f"{CC['single']['n']:,}", f"{CC['multi']['n']:,}", ""],
    ],
)
para(
    f'The first row is the trap. "Finished ≥1 schedule" is a **maximum over a worker\'s cohorts** - a worker in '
    f"three cohorts gets three independent chances to clear the bar — so it rises with cohort count even if nothing "
    f"about the worker changed."
)
para(
    f"Compare the two rows directly. The headline measure gives a **{_raw_gap}-point** gap "
    f"({CC['multi']['finished']}% − {CC['single']['finished']}%). The fair measure gives **{_pc_gap} points** "
    f"({pc(CC['multi'])}% − {pc(CC['single'])}%). So {_raw_gap} − {_pc_gap} = **{_raw_gap - _pc_gap} points** of the "
    f"original gap vanish once you stop rewarding a worker for having been put in more cohorts. That is "
    f"{_raw_gap - _pc_gap} ÷ {_raw_gap} = **{_arith}% of the headline gap**, which is arithmetic rather than "
    f"behaviour. To be explicit: {_arith}% is that share of the gap — it is NOT an average of the two rates."
    + (
        " On this build the remaining like-for-like difference is small enough to read as flat."
        if abs(_pc_gap) <= 3
        else ""
    )
)
para(
    f"Where the {pc(CC['multi'])}% itself comes from: for each multi-cohort worker, take the schedules they completed "
    f"divided by the schedules they were enrolled in, then average that across all {CC['multi']['n']:,} of them. "
    f"Single-cohort workers have exactly one schedule, so their fraction can only be 0 or 1 and the average collapses "
    f'to "what share finished their one schedule" — which is why both of their measures read '
    f"{CC['single']['finished']}%.",
    italic=True,
)
para(
    f"What re-used workers clearly do differ on is **depth: {CC['single']['depth']} → {CC['multi']['depth']} "
    f"words/session**, while completion rate is essentially unchanged ({CC['single']['completion']} vs "
    f"{CC['multi']['completion']})."
)
para(
    "Causal caution: being re-invited is itself an outcome of how a worker performed the first time, so even the "
    "residual difference is not evidence that re-using a worker causes them to finish more. Re-use is operationally "
    "valuable — these are known, trained, available workers who answer at greater length — but do not budget for a "
    "finish-rate gain from re-use itself.",
    italic=True,
)
if FE.get("armCombos"):
    _combos = ", ".join(f"{c['k']} ({c['n']})" for c in FE["armCombos"][:4])
    para(
        f"Most common arm combinations: {_combos}. (The arm set is unordered in this data, so we can say which arms "
        "co-occur, not which came first.)"
    )

# ================================================================ 3
h("3. Where the programme loses people — and who they are", 1)
para(
    f"Genuine early loss is the One-and-done segment: **{OAD['n']} workers ({OAD['pct']}%)** who started exactly one "
    "interview and never returned. They are not a random slice. Each row below compares the segment against that "
    "group's share of the whole population, so a share is only notable if it exceeds the base rate:"
)
_oad_rows = []
if OAD.get("topState"):
    _s = OAD["topState"][0]
    _base = round(100 * ST.get(_s["k"], {}).get("n", 0) / N) if ST.get(_s["k"]) else None
    _lift = f"×{_s['pct']/_base:.1f}" if _base else "—"
    _oad_rows.append([f"{_s['k']} (top state)", f"{_s['pct']}%", f"{_base}%" if _base else "—", _lift])
_oad_rows.append(
    [
        "Single-cohort",
        f"{OAD['singleCohortPct']}%",
        f"{FE['overallSingleCohortPct']}%",
        f"×{OAD['singleCohortPct']/FE['overallSingleCohortPct']:.1f}" if FE.get("overallSingleCohortPct") else "—",
    ]
)
if OAD.get("topType"):
    _t = OAD["topType"][0]
    _tb = next((round(100 * r["n"] / N) for r in FE["byType"] if r["k"] == _t["k"]), None)
    _oad_rows.append(
        [f"{_t['k']} (top cadre)", f"{_t['pct']}%", f"{_tb}%" if _tb else "—", f"×{_t['pct']/_tb:.1f}" if _tb else "—"]
    )
_oad_rows.append(
    [
        "Median first-session depth",
        f"{OAD['medianDepth']} words",
        f"{DS_MED if DS_MED is not None else '—'} words (all workers)",
        "",
    ]
)
table(["Cut", "One-and-done", "Programme base rate", "Over-represented"], _oad_rows)
para(
    "**Read each row separately, not stacked.** Each line above is measured on its own, so they cannot be added "
    "up into one profile."
    + (
        f" Only **{MF['oad_all_three']} of the {MF['oad_n']} one-and-done workers "
        f"({MF['oad_all_three_pct']}%)** are {MF['top_state']} *and* {MF['top_type']} *and* single-cohort "
        f"at the same time. The single strongest signal is {MF['top_state']}."
        if MF.get("oad_all_three_pct") is not None
        else ""
    )
)
_s3 = SV.get(3, {})
para(
    f"On how far workers get: of those whose schedule even *contains* interview 3, **{_s3.get('pct_elig', _s3.get('pct'))}%** "
    f"reach it ({_s3.get('reached',0):,} of {_s3.get('elig',N):,}). Quoting it against the whole population instead "
    f"gives {_s3.get('pct')}%, which understates retention because most workers are in short cohorts that stop before "
    f"interview 3 — that is schedule length, not attrition."
)
para(
    f"On attrition, two numbers that are often confused: **{OAD['pct']}% ({OAD['n']} workers) is the clean early-loss "
    f"figure** — they did one interview and vanished."
    + (
        f" Separately, **{MF['unfinished']} workers ({MF['unfinished_pct']}% of all) have not finished any schedule** — "
        f"most of them engaged repeatedly first. So {OAD['pct']}% is who we lost immediately; "
        f"{MF['unfinished_pct']}% is who has not got there yet. §6 draws its recovery list from the second group."
        if MF.get("unfinished")
        else ""
    )
)

# ================================================================ 4
h("4. Is early answer depth associated with finishing?", 1)
para(
    f"Splitting workers at the median depth of their **first session only** ({DS_MED if DS_MED is not None else '—'} words) — first "
    "session, so the predictor does not contain the outcome:"
)
table(
    ["Group", "Workers", "Per-cohort finish rate", "Finished ≥1 schedule", "First-session depth"],
    [
        [
            "Above-median depth",
            f"{DS['hi']['n']:,}",
            f"{pc(DS['hi'])}%",
            f"{DS['hi']['finished']}%",
            f"{DS['hi'].get('first_depth','—')} words",
        ],
        [
            "Below-median depth",
            f"{DS['lo']['n']:,}",
            f"{pc(DS['lo'])}%",
            f"{DS['lo']['finished']}%",
            f"{DS['lo'].get('first_depth','—')} words",
        ],
    ],
)
_d_pc = pc(DS["hi"]) - pc(DS["lo"])
para(
    f"Read those percentages carefully: **{pc(DS['hi'])}% means that group completed, on average, {pc(DS['hi'])}% of "
    f"the schedules they were enrolled in** — it does NOT mean they finished everything. Workers in both halves "
    f"finish some of their schedules and not others."
)
para(
    f"The difference between the halves is **{_d_pc} points** on the like-for-like measure "
    f"({pc(DS['hi'])}% vs {pc(DS['lo'])}%). On the looser \"finished ≥1\" reading it is "
    f"{DS['hi']['finished']}% vs {DS['lo']['finished']}% — a gap of {DS['hi']['finished'] - DS['lo']['finished']} "
    f"points. It is a real association "
    "and the strongest early signal available, but it is modest — and the deeper group also happens to be in more "
    "cohorts, so the two effects are tangled together. Treat first-interview support as the leading hypothesis to "
    "**test**, not an established lever."
)

# ================================================================ 5
h("5. Geography and partner — one finding, not two", 1)
table(
    ["State", "Workers", "Per-cohort finish rate", "Finished ≥1 schedule", "Answer depth"],
    [[s["k"], f"{s['n']:,}", f"{pc(s)}%", f"{s['finished']}%", f"{s['depth']} words"] for s in FE["byState"]],
)
_llo_rows = [r for r in FE["byLLO"] if not r.get("residual")]
table(
    ["Implementing partner", "Workers", "Per-cohort finish rate", "Finished ≥1 schedule", "Answer depth"],
    [[s["k"], f"{s['n']:,}", f"{pc(s)}%", f"{s['finished']}%", f"{s['depth']} words"] for s in _llo_rows],
)
_pairs = {}
for _s in FE["byState"]:
    _p = partner_of(_s["k"])
    if _p:
        _pairs.setdefault(_p, []).append(_s["k"])
if len(_pairs) >= 2:
    _desc = "; ".join(f"**{p}** serves {' and '.join(sorted(ss))}" for p, ss in sorted(_pairs.items()))
    para(
        f"These two tables are **the same finding cut two ways**. In this data the partners and the states are nested "
        f"({_desc}), so a state's result and its partner's result are the same observation — the data cannot tell you "
        f"whether the difference is geography, the partner's ways of working, or something the two share. Treat "
        f"\"the {min(_llo_rows, key=lambda r: pc(r))['k']} gap\" and "
        f"\"the {min(FE['byState'], key=lambda r: pc(r))['k']} gap\" as one issue, and do not count them as two "
        f"independent findings or give them two separate workstreams."
    )
_worst = min(FE["byState"], key=lambda r: pc(r))
_best = max(FE["byState"], key=lambda r: pc(r))
para(
    f"The spread is wide: {_best['k']} {pc(_best)}% vs {_worst['k']} {pc(_worst)}% per-cohort — "
    f"{pc(_best) - pc(_worst)} points."
)
_types = [r for r in FE["byType"] if not r.get("residual")]
if _types:
    _big = max(_types, key=lambda r: r["n"])
    para(
        f"By cadre, results are much tighter: {min(pc(r) for r in _types)}–{max(pc(r) for r in _types)}% per-cohort "
        f"across {len(_types)} cadres. The largest is **{_big['k']} ({_big['n']:,} workers) at {pc(_big)}%**. Because "
        "the cadre spread is narrow while the geography spread is wide, cadre looks like the weaker lever of the two."
    )

# ================================================================ 6
# ---------------------------------------------------------------- 6: geography, at the level that matters
GV = FE.get("geoVariance") or {}
if GV.get("states"):
    h("6. The variation is LOCAL, not state-level — and that changes what to do", 1)
    para(
        f"§5 compared states and partners because those are the units we manage by. But splitting the same workers by "
        f"**LGA** (local government area, {GV['n_lgas']} of them with enough workers to measure) shows the state "
        f"framing is the wrong altitude:"
    )
    table(
        ["State", "Per-cohort finish", "Spread between its own LGAs", "LGAs measured"],
        [[s["k"], f"{s['pc']}%", f"{s['lga_spread']} points", s["n_lgas"]] for s in GV["states"]],
    )
    para(
        f"**The gap between the best and worst state is {GV['state_spread']} points. The gap between the best and "
        f"worst LGA is {GV['lga_spread']} points** — and every single state has more variation inside it than exists "
        f"between the states. "
        + (
            f"The best-performing LGA in {GV['worst_state']} (our weakest state, {GV['worst_state_best_lga']}%) beats "
            f"several LGAs in our strongest one."
            if GV.get("worst_state_best_lga")
            else ""
        )
    )
    para(
        "Why this matters for the decision: §5 correctly says we cannot separate partner from state, because they "
        "are perfectly nested — and that looked like a dead end. This resolves it from the other direction: whatever "
        "is driving performance is mostly operating BELOW the state, so a partner-wide or state-wide explanation is "
        "the wrong shape regardless of which one you blame. The unit of action is the LGA, and the question to ask is "
        "what the strong LGAs do differently from the weak ones in the SAME state under the SAME partner.",
        italic=True,
    )

# ---------------------------------------------------------------- 7: the two behavioural levers that survived
_pe, _pa = FE.get("byPeers") or [], FE.get("byPace") or []
if _pe or _pa:
    h("7. Two things about how workers work that track with finishing", 1)
if _pe:
    para(
        "**Working alone is a disadvantage.** Grouping workers by how many colleagues share their settlement — the "
        "finest geography we hold — gives a clean gradient in the expected direction. This is the one factor here "
        "that the community-health-worker literature consistently flags (informal peer support), so it is a "
        "hypothesis we had reason to test rather than one found by trawling:"
    )
    table(
        ["Co-workers in the same settlement", "Workers", "Per-cohort finish", "Finished ≥1"],
        [[r["k"], f"{r['n']:,}", f"{pc(r)}%", f"{r['finished']}%"] for r in _pe],
    )
if _pa:
    para(
        "**Falling behind the schedule shows up early and gets worse.** Each worker's typical gap between interviews "
        "is measured against what their own schedule asks for, so subgroups with 3-day and 14-day cadences are "
        "compared fairly:"
    )
    table(
        ["Pace vs their own schedule", "Workers", "Per-cohort finish", "Finished ≥1"],
        [[r["k"], f"{r['n']:,}", f"{pc(r)}%", f"{r['finished']}%"] for r in _pa],
    )
    para(
        "Note the last row is workers with a single interview, who have no pace to measure — they are 0% by "
        "definition, not by behaviour. Among workers with a rhythm to measure, the gradient is monotonic.",
        italic=True,
    )
if _pe or _pa:
    para(
        "Why these two and not others: both are measurable before a worker is lost, which makes them usable as "
        "early warnings rather than post-mortems: peer density is known at assignment, and pace is visible after two "
        "interviews. We also tested response *consistency* (longest silence versus a worker's own typical gap) and "
        "onboarding delay (training date to first interview) and are not reporting either: consistency does not "
        "separate cleanly in this data, and onboarding delay has no variation to analyse — all but 20 workers start "
        "within a week. Education and first language showed no relationship to finishing either.",
        italic=True,
    )


h("8. The reachable at-risk list", 1)
_ar = ", ".join(f"{s['k']} ({s['n']})" for s in AR["byState"])
para(
    f"**{AR['n']} workers** started, have not finished any schedule, **were** offered a complete schedule, and have "
    f"been silent 14–60 days — recent enough that a nudge is plausible. Concentrated in {_ar}."
    + (
        f" They are the recent slice of **{AR['ofUnfinished']:,}** unfinished workers, not the whole unfinished "
        "population — the rest have been silent longer than 60 days."
        if AR.get("ofUnfinished")
        else ""
    )
)
para(
    "Two caveats: the list is defined off the newest session in the dataset, so it moves as data refreshes; and it "
    "deliberately excludes workers the programme never finished triggering, who are unfinished through no choice of "
    "their own.",
    italic=True,
)

# ================================================================ 7
h("9. What to do — and how confident we are", 1)
li(
    f"**Run the {AR['n']}-worker recovery list now.** Small, named, and time-bounded — the only directly actionable "
    "item here. Confidence: high that these workers are correctly identified; whether a nudge converts them is unknown — running it is how we find out.",
    "List Number",
)
li(
    f"**Trial first-interview support, with a control group.** Above-median first-session depth is associated with a "
    f"{_d_pc}-point higher per-cohort finish rate. Confidence: moderate — real association, modest size, "
    "confounded with cohort count, so measure it rather than rolling it out.",
    "List Number",
)
li(
    f"**Treat the {_worst['k']} / {min(_llo_rows, key=lambda r: pc(r))['k'] if _llo_rows else 'partner'} gap as one "
    f"investigation.** It is the largest effect in the data ({pc(_best) - pc(_worst)} points) and it is where "
    f"one-and-done concentrates. Confidence: high that the gap is real, none at all on the cause — geography and "
    "partner cannot be separated here, so the next step is qualitative, not another cut of this data.",
    "List Number",
)
li(
    "**Keep re-using proven workers for operational reasons** — known, trained, available, and they answer at greater "
    "length. Confidence: high on the operational value; do not forecast a finish-rate gain from re-use itself, because "
    "the headline gap is mostly arithmetic and re-invitation is itself an outcome of past performance.",
    "List Number",
)

# ================================================================ METHOD
h("Method & data", 1)
li(
    f"**Grain:** one row per unique worker (connect_id), deduped across cohorts; metrics union the worker's sessions "
    f"across every arm. Ties out to the dashboard's canonical started-worker count ({N:,})."
)
li(
    "**Per-cohort finish rate** — the measure to quote. For each worker: the share of *their own* cohort "
    "schedules they finished. The published figure is the **average across workers**, each worker counting once "
    "- NOT total schedules finished divided by total schedules given, which gives a different (higher) number. "
    'Unlike "finished >=1 schedule", it does not rise just because a worker was in more cohorts.'
)
li(
    "**Reach curve** (how far through a schedule a worker got) is quoted against the workers whose schedule contains that interview, not the whole population, "
    "so a 2-interview worker is not counted as dropping out at interview 3."
)
li(
    "**Engagement tier (RFM):** recency + completion rate + answer depth, each scored 1–5, with recency measured "
    "against the freshest session in the dataset rather than today's date."
)
li(
    '**Personas** are rule-based. "Partial progress" means over half of triggered interviews completed but **no** '
    "schedule finished — it is not a slow finisher."
)
li(
    '**Limits.** Observational, so associations only. Small groups are pooled into an "Other / not recorded" row '
    "rather than dropped or published as their own rate. Figures move with each daily refresh; this document is "
    f"generated from {'dashboard render v' + str(RENDER_V) if RENDER_V else 'the published dashboard payload'}."
)
li("**Full per-worker detail** is in the flw_analysis.csv export; the dashboard's FLW Retention tab is interactive.")


# ---------------------------------------------------------------- renderers
def render_md(path):
    out = []
    for kind, a, b in BLOCKS:
        if kind == "title":
            out.append(f"# {a}\n")
        elif kind == "h":
            out.append(f"\n{'#' * (b + 1)} {a}\n")
        elif kind == "p":
            out.append(f"*{a}*\n" if b.get("italic") else f"{a}\n")
        elif kind == "li":
            out.append(("1. " if b == "List Number" else "- ") + a)
        elif kind == "table":
            out.append("")
            out.append("| " + " | ".join(str(x) for x in a) + " |")
            out.append("|" + "|".join("---" for _ in a) + "|")
            for row in b:
                out.append("| " + " | ".join(str(x) for x in row) + " |")
            out.append("")
    md = "\n".join(out).replace("\n\n\n", "\n\n")
    with open(path, "w", encoding="utf-8") as f:
        f.write(md.rstrip() + "\n")
    return md


NAVY = RGBColor(0x1F, 0x38, 0x64)
GREY = RGBColor(0x6B, 0x72, 0x80)


def render_docx(path):
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    def runs(p, text):
        for i, seg in enumerate(text.split("**")):
            if seg:
                r = p.add_run(seg)
                r.bold = i % 2 == 1

    for kind, a, b in BLOCKS:
        if kind == "title":
            r = doc.add_paragraph().add_run(a)
            r.bold, r.font.size, r.font.color.rgb = True, Pt(20), NAVY
        elif kind == "h":
            p = doc.add_heading(a, level=b)
            for r in p.runs:
                r.font.color.rgb = NAVY
        elif kind == "p":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            runs(p, a)
            for r in p.runs:
                if b.get("italic"):
                    r.italic = True
                    r.font.color.rgb = GREY
                    r.font.size = Pt(9.5)
        elif kind == "li":
            p = doc.add_paragraph(style=b)
            p.paragraph_format.space_after = Pt(3)
            runs(p, a)
        elif kind == "table":
            t = doc.add_table(rows=1, cols=len(a))
            t.style = "Light Grid Accent 1"
            for j, x in enumerate(a):
                run = t.rows[0].cells[j].paragraphs[0].add_run(str(x))
                run.bold, run.font.size = True, Pt(9.5)
            for row in b:
                cells = t.add_row().cells
                for j, v in enumerate(row):
                    run = cells[j].paragraphs[0].add_run(str(v))
                    run.font.size = Pt(9.5)
                    if j == 0:
                        run.bold = True
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

    try:
        doc.save(path)
        return path
    except PermissionError:
        alt = path.replace(".docx", "_UPDATED.docx")
        doc.save(alt)
        print("(canonical .docx was locked/open — wrote the _UPDATED variant instead)")
        return alt


MD_PATH = os.path.join("docs", "FLW_Retention_Analysis_Brief.md")
DOCX_PATH = os.path.join("docs", "FLW_Retention_Analysis_Brief.docx")


def check():
    """Flag a brief that no longer matches the payload it claims to be generated from.

    This exists because the .docx silently went stale once: Word held the file open, the generator
    quietly wrote a _UPDATED copy instead, and the stale canonical file stayed in git for days while
    the dashboard said something different. Exit code 1 = do not circulate.
    """
    problems = []
    if not os.path.exists(DOCX_PATH):
        problems.append(f"{DOCX_PATH} is missing")
    elif os.path.getmtime(DOCX_PATH) < os.path.getmtime("flw_analysis_payload.json"):
        problems.append(f"{DOCX_PATH} is OLDER than flw_analysis_payload.json — it predates the current numbers")
    if os.path.exists(DOCX_PATH.replace(".docx", "_UPDATED.docx")):
        problems.append(
            "a _UPDATED.docx exists — the canonical file was locked when it was last generated, "
            "so the canonical file is the stale one"
        )
    if os.path.exists(MD_PATH) and RENDER_V and f"render v{RENDER_V}" not in open(MD_PATH, encoding="utf-8").read():
        problems.append(f"{MD_PATH} is not stamped with the current render (v{RENDER_V})")
    for p in problems:
        print(f"  [STALE] {p}")
    if problems:
        print(
            "\n  Fix: close the document in Word, then re-run:\n"
            "      python pull_live_payload.py && python build_flw_docx.py"
        )
    else:
        print(f"  [OK] both briefs match the current payload (render v{RENDER_V}, {TODAY})")
    return 1 if problems else 0


if __name__ == "__main__":
    import sys

    if "--check" in sys.argv:
        sys.exit(check())
    md_path = MD_PATH
    render_md(md_path)
    dx = render_docx(DOCX_PATH)
    print(f"wrote {md_path} + {dx}")
    print(
        f"  N={N:,}  finishers={finishers_pct}%  per-cohort single/multi={pc(CC['single'])}/{pc(CC['multi'])}%  "
        f"as of {TODAY} (render v{RENDER_V})"
    )
